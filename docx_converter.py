import os
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor
import zipfile
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class MarkdownToDocxConverter:
    def __init__(self, input_path, output_folder):
        self.input_path = input_path
        self.output_folder = output_folder
        self.document = None
        self.converted_files = []
        try:
            os.makedirs(self.output_folder, exist_ok=True)
        except Exception as e:
            raise RuntimeError(f"Failed to create output directory: {e}")

    def _embed_font(self, font_name):
        """Embed font in the document."""
        embed = OxmlElement('w:embedRegular')
        embed.set(qn('w:fontKey'), f"{font_name}_Regular")
        embed.set(qn('w:subsetted'), "1")
        return embed

    def _set_document_styles(self):
        section = self.document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.orientation = WD_ORIENT.PORTRAIT
        section.different_first_page_header_footer = True

        # Set up normal style with font embedding
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Helvetica'
        font.size = Pt(11)
        
        # Add fallback fonts
        rFonts = font._element.rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Helvetica')
        rFonts.set(qn('w:hAnsi'), 'Helvetica')
        rFonts.set(qn('w:cs'), 'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Embed fonts
        embed_regular = self._embed_font('Helvetica')
        font._element.rPr.append(embed_regular)

        # Set up heading styles with font embedding
        for i in range(1, 10):
            heading_style = self.document.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'Helvetica'
            
            # Add fallback fonts for headings
            rFonts = heading_font._element.rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Helvetica')
            rFonts.set(qn('w:hAnsi'), 'Helvetica')
            rFonts.set(qn('w:cs'), 'Arial')
            rFonts.set(qn('w:eastAsia'), 'Arial')
            
            # Embed fonts for headings
            embed_regular = self._embed_font('Helvetica')
            heading_font._element.rPr.append(embed_regular)

    def convert(self):
        if os.path.isfile(self.input_path):
            self._convert_file(self.input_path)
        elif os.path.isdir(self.input_path):
            for filename in os.listdir(self.input_path):
                if filename.endswith(('.md', '.txt')):
                    file_path = os.path.join(self.input_path, filename)
                    self._convert_file(file_path)
        else:
            raise ValueError("Input path is neither a file nor a directory")
        self._create_zip()

    def _convert_file(self, file_path):
        self.document = Document()
        self._set_document_styles()
        output_filename = os.path.splitext(os.path.basename(file_path))[0] + '.docx'
        output_path = os.path.join(self.output_folder, output_filename)
        with open(file_path, 'r', encoding='utf-8') as md_file:
            md_content = md_file.read()
        html = markdown.markdown(md_content, extensions=['extra'])
        self._add_html_to_document(html)
        self.document.save(output_path)
        self.converted_files.append(output_path)

    def _add_html_to_document(self, html):
        soup = BeautifulSoup(html, 'html.parser')
        for element in soup.find_all():
            if element.name == 'p':
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._set_first_line_indent(paragraph)
                self._process_inline_elements(element, paragraph)
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                paragraph = self.document.add_paragraph(style=f'Heading {level}')
                paragraph.text = element.get_text()
            elif element.name == 'ul':
                self._add_list(element, is_numbered=False)
            elif element.name == 'ol':
                self._add_list(element, is_numbered=True)
            elif element.name == 'blockquote':
                paragraph = self.document.add_paragraph(style='Quote')
                self._process_inline_elements(element, paragraph)

    def _process_inline_elements(self, element, paragraph):
        for node in element.contents:
            if isinstance(node, str):
                # Direct text content
                if node.strip():
                    paragraph.add_run(node)
            elif node.name is None and node.string:
                # NavigableString
                if node.string.strip():
                    paragraph.add_run(node.string)
            elif node.name in ['strong', 'b']:
                run = paragraph.add_run(node.get_text())
                run.bold = True
            elif node.name in ['em', 'i']:
                run = paragraph.add_run(node.get_text())
                run.italic = True
            elif node.name == 'u':
                run = paragraph.add_run(node.get_text())
                run.underline = True
            elif node.name == 'a':
                run = paragraph.add_run(node.get_text())
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)

    def _add_list(self, list_element, is_numbered):
        for item in list_element.find_all('li', recursive=False):
            style = 'List Number' if is_numbered else 'List Bullet'
            paragraph = self.document.add_paragraph(style=style)
            
            # Process all content of the list item
            for node in item.contents:
                if isinstance(node, str):
                    # Direct text content
                    if node.strip():
                        paragraph.add_run(node)
                elif node.name is None and node.string:
                    # NavigableString
                    if node.string.strip():
                        paragraph.add_run(node.string)
                elif node.name == 'p':
                    # Paragraph inside list item
                    self._process_inline_elements(node, paragraph)
                elif node.name in ['strong', 'b']:
                    run = paragraph.add_run(node.get_text())
                    run.bold = True
                elif node.name in ['em', 'i']:
                    run = paragraph.add_run(node.get_text())
                    run.italic = True
                elif node.name == 'u':
                    run = paragraph.add_run(node.get_text())
                    run.underline = True
                elif node.name == 'a':
                    run = paragraph.add_run(node.get_text())
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif node.name in ['ul', 'ol']:
                    # Nested list
                    self._add_list(node, is_numbered=(node.name == 'ol'))

    def _set_first_line_indent(self, paragraph):
        p_fmt = paragraph.paragraph_format
        p_fmt.first_line_indent = Cm(1.27)

    def _create_zip(self):
        zip_filename = os.path.basename(os.path.join(self.output_folder, os.path.basename(self.output_folder) + '.zip'))
        zip_path = os.path.join(self.output_folder, zip_filename)
        os.makedirs(os.path.dirname(zip_path), exist_ok=True)
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in self.converted_files:
                zipf.write(file, os.path.basename(file))
